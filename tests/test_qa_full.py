"""Full-system QA / regression: real end-to-end CRUD, backups and agent tools.

Every test here performs ACTUAL data actions through the running Flask app
(test client = real HTTP round-trips against a temp SQLite DB). Nothing is
mocked except the LLM network calls (agent decisions), exactly like the rest
of the unit suite. Run with:  python -m unittest discover tests
"""
import io
import json
import os
import sqlite3
import sys
import tempfile
import unittest
import urllib.parse
from datetime import date
from pathlib import Path
from unittest import mock

BASE_DIR = Path(__file__).resolve().parent.parent
_ENV_ROOT = Path(tempfile.mkdtemp(prefix="assistant-qa-"))
os.environ["ASSISTANT_DB"] = str(_ENV_ROOT / "qa.db")
os.environ["ASSISTANT_UPLOADS"] = str(_ENV_ROOT / "uploads")

sys.path.insert(0, str(BASE_DIR))
import app as app_module  # noqa: E402


class QaBase(unittest.TestCase):
    """Sandboxed app + auth, mirrors the suite BaseTest without importing it."""

    def setUp(self):
        self.app = app_module.app
        self.client = self.app.test_client()
        # No network for any chat/agent test: strip provider config first.
        self._genv_key = os.environ.pop("GEMINI_API_KEY", None)
        self._genv_model = os.environ.pop("GEMINI_MODEL", None)
        for p in tuple(app_module.CHAT_PROVIDERS):
            if p == "gemini":
                continue
            os.environ.pop(f"{p.upper()}_API_KEY", None)
            os.environ.pop(f"{p.upper()}_MODEL", None)
        app_module.reset_db()
        conn = sqlite3.connect(app_module.DB_PATH)
        conn.executescript(
            "DELETE FROM users; DELETE FROM app_settings; DELETE FROM agent_pending; "
            "DELETE FROM chat_agents; DELETE FROM api_tools;"
        )
        conn.commit()
        conn.close()
        self.register()
        self.login()

    def tearDown(self):
        for p in tuple(app_module.CHAT_PROVIDERS):
            os.environ.pop(f"{p.upper()}_API_KEY", None)
            os.environ.pop(f"{p.upper()}_MODEL", None)
        if self._genv_key is not None:
            os.environ["GEMINI_API_KEY"] = self._genv_key
        if self._genv_model is not None:
            os.environ["GEMINI_MODEL"] = self._genv_model
        app_module._CHAT_FILE_BUCKET.clear()

    def register(self, username="tester", password="secret1", display_name="Tester"):
        return self.client.post(
            "/api/auth/register",
            json={"username": username, "password": password, "display_name": display_name},
        )

    def login(self, username="tester", password="secret1"):
        return self.client.post("/api/auth/login", json={"username": username, "password": password})

    def _fresh_session(self):
        return self.client.post("/api/chat/sessions", json={}).get_json()["id"]

    def _set(self, key, value):
        app_module._set_app_setting(key, str(value))


# ================================ 1. FULL CRUD ================================


class TestQaCrud(QaBase):
    """Tasks / Notes / Pages / Routines / Knowledge: create, read, update, delete."""

    def test_task_full_crud(self):
        # CREATE
        r = self.client.post("/api/tasks", json={
            "title": "QA Task A", "description": "verify auth",
            "priority": "high", "due_date": "2026-09-15",
        })
        self.assertEqual(r.status_code, 201)
        tid = r.get_json()["id"]
        self.assertEqual(r.get_json()["done"], 0)
        # READ (dashboard list reflects immediately)
        titles = [t["title"] for t in self.client.get("/api/tasks").get_json()]
        self.assertIn("QA Task A", titles)
        # UPDATE status + due date
        r = self.client.patch(f"/api/tasks/{tid}", json={"done": True, "due_date": "2026-09-20"})
        self.assertEqual(r.status_code, 200)
        got = next(t for t in self.client.get("/api/tasks").get_json() if t["id"] == tid)
        self.assertEqual(got["done"], 1)
        self.assertEqual(got["due_date"], "2026-09-20")
        self.assertTrue(got["completed_at"])
        # DELETE
        self.assertEqual(self.client.delete(f"/api/tasks/{tid}").status_code, 200)
        titles = [t["title"] for t in self.client.get("/api/tasks").get_json()]
        self.assertNotIn("QA Task A", titles)

    def test_note_full_crud_with_versions_and_restore(self):
        # CREATE with formatting
        r = self.client.post("/api/notes", json={
            "title": "QA Note", "tags": "qa,billing",
            "content": "<p><strong>Bold</strong> <em>italic</em> text</p>",
        })
        self.assertEqual(r.status_code, 201)
        nid = r.get_json()["id"]
        notes = self.client.get("/api/notes").get_json()
        self.assertIn("QA Note", [n["title"] for n in notes])
        seen = next(n["content"] for n in notes if n["id"] == nid)
        self.assertIn("<strong>Bold</strong>", seen)  # sanitizer keeps safe formatting
        # UPDATE -> the pre-change content is snapshotted as a new version
        self.assertEqual(self.client.put(f"/api/notes/{nid}", json={"content": "<p>Version two content</p>"}).status_code, 200)
        self.assertEqual(self.client.put(f"/api/notes/{nid}", json={"content": "<p>Version three content</p>"}).status_code, 200)
        versions = self.client.get(f"/api/notes/{nid}/versions").get_json()
        self.assertEqual(len(versions), 2)
        oldest = versions[-1]  # the ORIGINAL "Bold" snapshot
        snapshot = self.client.get(f"/api/notes/{nid}/versions/{oldest['id']}").get_json()
        self.assertIn("Bold", snapshot["content"])
        # RESTORE the original snapshot
        r = self.client.post(f"/api/notes/{nid}/restore", json={"version_id": oldest["id"]})
        self.assertEqual(r.status_code, 200)
        restored = next(n for n in self.client.get("/api/notes").get_json() if n["id"] == nid)
        self.assertIn("Bold", restored["content"])
        # DELETE
        self.assertEqual(self.client.delete(f"/api/notes/{nid}").status_code, 200)
        self.assertNotIn("QA Note", [n["title"] for n in self.client.get("/api/notes").get_json()])

    def test_page_full_crud(self):
        r = self.client.post("/api/pages", json={"title": "QA Page", "icon": "*", "content": "<p>Page body</p>"})
        self.assertEqual(r.status_code, 201)
        pid = r.get_json()["id"]
        self.assertIn("QA Page", [p["title"] for p in self.client.get("/api/pages").get_json()])
        r = self.client.put(f"/api/pages/{pid}", json={"title": "QA Page v2", "content": "<p>Updated body</p>"})
        self.assertEqual(r.status_code, 200)
        self.assertIn("QA Page v2", [p["title"] for p in self.client.get("/api/pages").get_json()])
        self.assertIn("Updated body", body := self.client.get("/api/pages").get_json()[0]["content"])
        self.assertEqual(self.client.delete(f"/api/pages/{pid}").status_code, 200)
        self.assertNotIn("QA Page", [p["title"] for p in self.client.get("/api/pages").get_json()])

    def test_routine_full_crud_and_agent_fetch(self):
        # CREATE
        r = self.client.post("/api/routines", json={"title": "QA Routine", "weekday": 2, "time": "10:00"})
        self.assertEqual(r.status_code, 201)
        rid = r.get_json()["id"]
        self.assertEqual(r.get_json()["completions"], [])
        self.assertIn("QA Routine", [x["title"] for x in self.client.get("/api/routines").get_json()])
        # UPDATE
        r = self.client.patch(f"/api/routines/{rid}", json={"time": "11:30"})
        self.assertEqual(r.status_code, 200)
        got = next(x for x in self.client.get("/api/routines").get_json() if x["id"] == rid)
        self.assertEqual(got["time"], "11:30")
        # COMPLETE (toggle)
        r = self.client.post(f"/api/routines/{rid}/toggle", json={"date": date.today().isoformat()})
        self.assertEqual(r.get_json()["done"], True)
        got = next(x for x in self.client.get("/api/routines").get_json() if x["id"] == rid)
        self.assertIn(date.today().isoformat(), got["completions"])
        # ADMINISTRATOR AGENT fetches the routine through the real SQL tool
        self._set("agent_enabled", "1")
        self._set("live_chat_ai", "0")
        self._set("review_enabled", "0")
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        try:
            side = [{"action": "sql", "kind": "sql", "id": None, "title": "", "fields": {},
                     "query": "SELECT title, weekday, time FROM routines ORDER BY id DESC"}]
            with mock.patch.object(app_module, "_agent_plan", return_value=side[0]), \
                 mock.patch.object(app_module, "_llm_prompt", return_value="Fetched the routine for you."):
                sid = self._fresh_session()
                text, source, outcome = app_module.agent_answer(
                    sid, "routines batao", "Tester", False)
            self.assertEqual(outcome, "action")
            self.assertEqual(source, "agent_action")
            # The raw tool output really contains the routine (proxy for the fetch succeeding)
            plan = side[0]
            raw, ok = app_module._run_readonly_sql(plan["query"])
            self.assertTrue(ok)
            self.assertIn("QA Routine", raw)
        finally:
            os.environ.pop("GEMINI_API_KEY", None)
        # DELETE
        self.assertEqual(self.client.delete(f"/api/routines/{rid}").status_code, 200)
        self.assertNotIn("QA Routine", [x["title"] for x in self.client.get("/api/routines").get_json()])

    def test_knowledge_full_crud(self):
        r = self.client.post("/api/knowledge", json={"title": "QA Rule", "category": "CPT", "content": "Submit in 14 days"})
        self.assertEqual(r.status_code, 201)
        kid = r.get_json()["id"]
        self.assertIn("QA Rule", [k["title"] for k in self.client.get("/api/knowledge").get_json()])
        r = self.client.put(f"/api/knowledge/{kid}", json={"content": "Submit in 30 days"})
        self.assertEqual(r.status_code, 200)
        got = next(k for k in self.client.get("/api/knowledge").get_json() if k["id"] == kid)
        self.assertIn("30 days", got["content"])
        self.assertEqual(self.client.delete(f"/api/knowledge/{kid}").status_code, 200)
        self.assertNotIn("QA Rule", [k["title"] for k in self.client.get("/api/knowledge").get_json()])

    def test_dashboard_endpoints_reflect_changes_without_refresh(self):
        a = self.client.post("/api/tasks", json={"title": "Dash Task"}).get_json()["id"]
        self.client.post("/api/routines", json={"title": "Dash Routine", "weekday": 1, "time": "08:00"})
        # The dashboard is a composition of these endpoints; each must re-shape instantly.
        self.assertIn("Dash Task", [t["title"] for t in self.client.get("/api/tasks").get_json()])
        self.assertIn("Dash Routine", [x["title"] for x in self.client.get("/api/routines").get_json()])
        self.client.patch(f"/api/tasks/{a}", json={"done": True})
        got = next(t for t in self.client.get("/api/tasks").get_json() if t["id"] == a)
        self.assertEqual(got["done"], 1)
        for path in ("/api/notes", "/api/pages", "/api/knowledge", "/api/routines", "/api/tasks"):
            self.assertEqual(self.client.get(path).status_code, 200)


# ============================ 2. BACKUP INTEGRITY ============================


class TestQaBackup(QaBase):
    """Exports captured the latest schema; dummy restore boots cleanly."""

    def setUp(self):
        super().setUp()
        self.client.post("/api/tasks", json={"title": "Backup Task", "priority": "high", "due_date": "2026-09-10"})
        self.client.post("/api/notes", json={"title": "Backup Note", "content": "<p>pre-restore</p>"})

    def test_json_export_includes_recent_data(self):
        r = self.client.get("/api/export/json")
        self.assertEqual(r.status_code, 200)
        payload = json.loads(r.get_data())
        self.assertIn("data", payload)
        self.assertEqual(set(payload["data"]), set(app_module.BACKUP_TABLES))
        tasks = payload["data"]["tasks"]
        self.assertEqual(len(tasks), 1)  # Backup Task created in setUp
        self.assertIn("Backup Task", [t["title"] for t in tasks])
        notes = payload["data"]["notes"]
        self.assertEqual(notes[0]["title"], "Backup Note")

    def test_excel_export_has_all_sheets_and_rows(self):
        from openpyxl import load_workbook

        r = self.client.get("/api/export/excel")
        self.assertEqual(r.status_code, 200)
        wb = load_workbook(io.BytesIO(r.get_data()), read_only=True)
        self.assertTrue({"tasks", "notes", "pages", "routines", "routine_completions", "note_versions", "note_shares"} <= set(wb.sheetnames))
        rows = list(wb["tasks"].iter_rows(values_only=True))
        self.assertTrue(any("Backup Task" in (str(c) if c else "") for row in rows for c in row))

    def test_sqlite_export_has_latest_schema(self):
        r = self.client.get("/api/export/sqlite")
        self.assertEqual(r.status_code, 200)
        fd, tmp = tempfile.mkstemp(suffix=".sqlite")
        os.close(fd)
        with open(tmp, "wb") as fh:
            fh.write(r.get_data())
        conn = sqlite3.connect(tmp)
        tables = {row[0] for row in conn.execute("SELECT name FROM sqlite_master WHERE type='table'")}
        cols = {t: {c[1] for c in conn.execute(f"PRAGMA table_info({t})")} for t in ("tasks", "chat_agents", "agent_pending", "chat_messages")}
        rows = conn.execute("SELECT title FROM tasks").fetchall()
        conn.close()
        os.unlink(tmp)
        want = {"tasks", "notes", "pages", "routines", "routine_completions", "note_versions",
                "note_shares", "knowledge_base", "chat_sessions", "chat_messages", "chat_settings",
                "app_settings", "agent_pending", "chat_api_keys", "chat_agents", "api_tools",
                "embed_vectors", "users"}
        self.assertTrue(want <= tables, f"missing: {want - tables}")
        self.assertIn("page_id", cols["tasks"])
        self.assertIn("icon", cols["chat_agents"])
        self.assertIn("session_id", cols["agent_pending"])
        self.assertIn("source_type", cols["chat_messages"])
        self.assertTrue(any(r[0] == "Backup Task" for r in rows))

    def test_sqlite_dummy_restore_replaces_db_and_boots(self):
        # Snapshot the DB, then mutate it, then import the snapshot back.
        snap = self.client.get("/api/export/sqlite").get_data()
        self.client.post("/api/tasks", json={"title": "Post-backup task"})
        self.assertTrue(any(t["title"] == "Post-backup task" for t in self.client.get("/api/tasks").get_json()))
        # Checkpoint so WAL can't replay the newer state over the restored file.
        c = app_module.get_db()
        try:
            c.execute("PRAGMA wal_checkpoint(TRUNCATE)")
        except sqlite3.Error:
            pass
        c.close()
        r = self.client.post(
            "/api/import/sqlite",
            data={"file": (io.BytesIO(snap), "assistant-backup.sqlite")},
            content_type="multipart/form-data",
        )
        self.assertEqual(r.status_code, 200, r.get_json())
        self.assertEqual(r.get_json(), {"ok": True})
        # Data reverted to the snapshot; the system boots and answers afterwards.
        tasks = self.client.get("/api/tasks").get_json()
        self.assertTrue(any(t["title"] == "Backup Task" for t in tasks))
        self.assertFalse(any(t["title"] == "Post-backup task" for t in tasks))
        self.assertIn("Backup Note", [n["title"] for n in self.client.get("/api/notes").get_json()])
        # Fresh chat session works after restore (schema/boot intact)
        r = self.client.post("/api/chat/sessions", json={})
        self.assertEqual(r.status_code, 201)
        self.assertEqual(self.client.get("/api/auth/me").status_code, 200)

    def test_import_sqlite_rejects_garbage(self):
        r = self.client.post(
            "/api/import/sqlite",
            data={"file": (io.BytesIO(b"not a database at all"), "backup.sqlite")},
            content_type="multipart/form-data",
        )
        self.assertEqual(r.status_code, 400)
        self.client.post("/api/tasks", json={"title": "Still alive"})
        self.assertIn("Still alive", [t["title"] for t in self.client.get("/api/tasks").get_json()])


# ========================= 3. AGENT TOOLS + MAKER-CHECKER =========================


class TestQaAgentTools(QaBase):
    """Agents fetch data with the read-only SQL tool; maker-checker loop works."""

    def setUp(self):
        super().setUp()
        self._set("agent_enabled", "1")
        self._set("live_chat_ai", "0")
        self._set("review_enabled", "1")
        self._set("review_max_loops", "2")
        conn = sqlite3.connect(app_module.DB_PATH)
        conn.executescript("""
        INSERT INTO chat_agents (name, description, system_prompt, is_active, created_at, icon) VALUES
        ('Rumman Lashari', 'Administrator & Agent Coordinator - poore system ka boss.',
         'Aap Administrator hain.', 1, '2026-01-01 00:00:00', ''),
        ('Medical Billing', 'Medical Billing specialist - answers from Notes, Tasks & Guidelines',
         'Aap Medical Billing (RCM) specialist hain.', 1, '2026-01-01 00:00:00', '');
        """)
        conn.commit()
        conn.close()

    def test_agent_sql_fetch_runs_without_errors(self):
        self.client.post("/api/tasks", json={"title": "Fetch Me Task", "priority": "low"})
        self.client.post("/api/tasks", json={"title": "Fetch Me Task 2", "priority": "medium"})
        sid = self._fresh_session()
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        try:
            plan = {"action": "sql", "kind": "sql", "id": None, "title": "", "fields": {},
                    "query": "SELECT title, priority FROM tasks ORDER BY id DESC"}
            with mock.patch.object(app_module, "_agent_plan", return_value=plan), \
                 mock.patch.object(app_module, "_llm_prompt", return_value="Displayed 2 tasks."):
                text, source, outcome = app_module.agent_answer(sid, "tasks dikhao", "Tester", False)
            self.assertEqual(outcome, "action")
            self.assertEqual(source, "agent_action")
            self.assertTrue(text.startswith("Displayed 2 tasks."))  # attribution footer may follow
            # The real tool returned real rows (proxy: raw SQL execution)
            raw, ok = app_module._run_readonly_sql(plan["query"])
            self.assertTrue(ok)
            self.assertIn("Fetch Me Task", raw)
            self.assertNotIn("users", raw.lower())
        finally:
            os.environ.pop("GEMINI_API_KEY", None)

    def test_agent_users_table_refusal_is_graceful(self):
        """The users-table issue: querying it must NOT produce a 500 / raw error."""
        self.client.post("/api/pages", json={"title": "Staff Directory",
                                             "content": "Ali Raza - Gmail: ali.raza@example.com"})
        sid = self._fresh_session()
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        try:
            with mock.patch.object(
                app_module, "_agent_plan",
                return_value={"action": "sql", "kind": "sql", "id": None, "title": "", "fields": {},
                              "query": "SELECT email FROM users WHERE name = 'Ali'"},
            ), mock.patch.object(app_module, "_review_draft", return_value=("approved", "")):
                text, source, outcome = app_module.agent_answer(sid, "Ali ka email kya hai", "Tester", False)
            self.assertNotIn("I couldn't do that", text)
            self.assertNotIn("allow nahi hai", text)
            self.assertIn("email", text.lower()) if "email" in text.lower() else self.assertIn("ali", text.lower())
            self.assertIn("ali.raza", text.lower())
        finally:
            os.environ.pop("GEMINI_API_KEY", None)

    def test_maker_checker_loop_rejects_then_approves(self):
        """A plan rejected by the reviewer is corrected by the worker and approved."""
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        plan0 = {"action": "create", "kind": "task", "id": None, "title": "",
                 "fields": {"title": "Follow up claim", "description": "", "priority": "medium", "due_date": ""}}
        corrected = {"action": "create", "kind": "task", "id": None, "title": "",
                     "fields": {"title": "Follow up claim", "description": "Amount: 300",
                                "priority": "medium", "due_date": "2026-09-05"}}
        try:
            def fake_review(*a, **k):
                calls.append(k)
                return ("rejected", "amount figuer missing") if len(calls) == 1 else ("approved", "")

            calls = []
            with mock.patch.object(app_module, "_agent_plan", return_value=plan0), \
                 mock.patch.object(app_module, "_review_draft", side_effect=fake_review), \
                 mock.patch.object(app_module, "_llm_prompt", return_value=json.dumps(corrected)):
                text, source, outcome = app_module.agent_answer("sess-mc", "add task follow up claim", "Tester", False)
            self.assertEqual(outcome, "action")
            self.assertEqual(len(calls), 2)  # reviewer loop iterated: reject -> approve
        finally:
            os.environ.pop("GEMINI_API_KEY", None)
        tasks = self.client.get("/api/tasks").get_json()
        self.assertEqual(len(tasks), 1)
        self.assertEqual(tasks[0]["title"], "Follow up claim")
        self.assertEqual(tasks[0]["description"], "Amount: 300")
        self.assertEqual(tasks[0]["due_date"], "2026-09-05")

    def test_maker_checker_manual_flag_when_loops_exhausted(self):
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        self._set("review_max_loops", "0")
        plan0 = {"action": "create", "kind": "task", "id": None, "title": "",
                 "fields": {"title": "Manual review task"}}
        try:
            with mock.patch.object(app_module, "_agent_plan", return_value=plan0), \
                 mock.patch.object(app_module, "_review_draft", return_value=("rejected", "bad figure")):
                text, source, outcome = app_module.agent_answer("sess-m2", "create a task", "Tester", False)
            self.assertEqual(outcome, "error")
            self.assertIn("EXECUTE nahi hua", text)
        finally:
            os.environ.pop("GEMINI_API_KEY", None)
        # Nothing was created because the action was blocked
        self.assertEqual(self.client.get("/api/tasks").get_json(), [])


class TestQaAgentMemory(QaBase):
    """agent_memory: schema, API CRUD, chat-driven SQL writes, refusal guards."""

    def setUp(self):
        super().setUp()
        self._set("agent_enabled", "1")
        self._set("live_chat_ai", "0")
        self._set("review_enabled", "1")
        conn = sqlite3.connect(app_module.DB_PATH)
        conn.executescript("""
        INSERT INTO chat_agents (name, description, system_prompt, is_active, created_at, icon) VALUES
        ('Rumman Lashari', 'Administrator & Agent Coordinator - poore system ka boss.',
         'Aap Administrator hain.', 1, '2026-01-01 00:00:00', ''),
        ('Medical Billing', 'Medical Billing specialist - answers from Notes, Tasks & Guidelines',
         'Aap Medical Billing (RCM) specialist hain.', 0, '2026-01-01 00:00:00', '');
        """)
        conn.commit()
        conn.close()

    def _agent_row(self, name="Rumman Lashari"):
        conn = sqlite3.connect(app_module.DB_PATH)
        conn.row_factory = sqlite3.Row
        row = conn.execute("SELECT * FROM chat_agents WHERE name = ?", (name,)).fetchone()
        conn.close()
        return row

    def test_schema_and_list_include_memory(self):
        conn = sqlite3.connect(app_module.DB_PATH)
        tables = {r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'")}
        cols = {c[1] for c in conn.execute("PRAGMA table_info(agent_memory)")}
        conn.close()
        self.assertIn("agent_memory", tables)
        self.assertTrue({"agent_id", "kind", "key", "content", "source", "created_by"} <= cols)
        aid = self._agent_row()["id"]
        r = self.client.post(f"/api/agents/{aid}/memory", json={"kind": "fact", "key": "fav_tea", "content": "Chai"})
        self.assertEqual(r.status_code, 201)
        ag = next(a for a in self.client.get("/api/agents").get_json()["agents"] if a["id"] == aid)
        self.assertEqual(len(ag["memory"]), 1)
        self.assertEqual(ag["memory"][0]["key"], "fav_tea")

    def test_memory_crud_endpoints(self):
        aid = self._agent_row()["id"]
        r = self.client.post(f"/api/agents/{aid}/memory", json={"kind": "preference", "key": "lang", "content": "Urdu"})
        self.assertEqual(r.status_code, 201)
        mid = r.get_json()["id"]
        self.assertEqual(
            self.client.put(f"/api/agents/{aid}/memory/{mid}",
                            json={"kind": "preference", "key": "lang", "content": "English"}).get_json()["content"],
            "English",
        )
        self.assertEqual(self.client.delete(f"/api/agents/{aid}/memory/{mid}").status_code, 200)
        ag = next(a for a in self.client.get("/api/agents").get_json()["agents"] if a["id"] == aid)
        self.assertEqual(ag["memory"], [])
        self.assertEqual(self.client.post("/api/agents/9999/memory", json={"content": "x"}).status_code, 404)

    def test_write_tool_insert_via_chat(self):
        aid = self._agent_row()["id"]
        sid = self._fresh_session()
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        plan = {"action": "sql", "kind": "sql", "id": None, "title": "", "fields": {},
                "query": "INSERT INTO agent_memory (agent_id, kind, key, content, source, created_by) "
                         "VALUES ((SELECT id FROM chat_agents WHERE name='Rumman Lashari'), "
                         "'fact', 'fav_tea', 'Chai', 'chat', 'Assistant')"}
        try:
            with mock.patch.object(app_module, "_agent_plan", return_value=plan), \
                 mock.patch.object(app_module, "_review_draft", return_value=("approved", "")), \
                 mock.patch.object(app_module, "_llm_prompt", return_value="Yaad rakh liya: Chai."):
                text, source, outcome = app_module.agent_answer(sid, "yaad rakh lo woh chai pasand karta hai", "Tester", False)
            self.assertEqual(outcome, "action")
        finally:
            os.environ.pop("GEMINI_API_KEY", None)
        conn = sqlite3.connect(app_module.DB_PATH)
        row = conn.execute("SELECT key, content FROM agent_memory WHERE agent_id = ?", (aid,)).fetchone()
        conn.close()
        self.assertIsNotNone(row)
        self.assertEqual(tuple(row), ("fav_tea", "Chai"))
        merged = app_module._agent_prompt_with_memory({"id": aid, "system_prompt": "Aap Administrator hain."})
        self.assertIn("Chai", merged)

    def test_write_tool_refuses_foreign_table_and_nonedml(self):
        sid = self._fresh_session()
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        bad = {"action": "sql", "kind": "sql", "id": None, "title": "", "fields": {},
               "query": "INSERT INTO tasks (title) VALUES ('hacked')"}
        try:
            with mock.patch.object(app_module, "_agent_plan", return_value=bad), \
                 mock.patch.object(app_module, "_review_draft", return_value=("approved", "")), \
                 mock.patch.object(app_module, "_llm_prompt", return_value="idhar koi task nahi"):
                text, source, outcome = app_module.agent_answer(sid, "task add karo", "Tester", False)
            self.assertNotIn("I couldn't do that", text)
        finally:
            os.environ.pop("GEMINI_API_KEY", None)
        conn = sqlite3.connect(app_module.DB_PATH)
        n = conn.execute("SELECT COUNT(*) FROM tasks").fetchone()[0]
        conn.close()
        self.assertEqual(n, 0)
        aid = self._agent_row()["id"]
        with self.assertRaises(ValueError):
            app_module._run_agent_write_sql("UPDATE agent_memory SET content='x'")
        with self.assertRaises(ValueError):
            app_module._run_agent_write_sql("INSERT INTO tasks (title) VALUES ('x')")
        with self.assertRaises(ValueError):
            app_module._run_agent_write_sql("DROP TABLE agent_memory")
        with self.assertRaises(ValueError):
            app_module._run_agent_write_sql(
                "INSERT INTO agent_memory (agent_id, content) VALUES (1, 'a'); "
                "INSERT INTO agent_memory (agent_id, content) VALUES (2, 'b')"
            )
        text, ok = app_module._run_agent_write_sql(
            "UPDATE agent_memory SET content='Updated' WHERE agent_id=" + str(aid)
        )
        self.assertTrue(ok)

    def test_agents_delete_clears_memory(self):
        aid = self._agent_row("Medical Billing")["id"]
        self.client.post(f"/api/agents/{aid}/memory",
                         json={"kind": "instruction", "key": "rule", "content": "Always verify CPT"})
        self.assertEqual(self.client.delete(f"/api/agents/{aid}").status_code, 200)
        conn = sqlite3.connect(app_module.DB_PATH)
        n = conn.execute("SELECT COUNT(*) FROM agent_memory WHERE agent_id = ?", (aid,)).fetchone()[0]
        still = conn.execute("SELECT COUNT(*) FROM chat_agents WHERE id = ?", (aid,)).fetchone()[0]
        conn.close()
        self.assertEqual(n, 0)
        self.assertEqual(still, 0)

    def test_readonly_sql_allows_memory_reads(self):
        aid = self._agent_row()["id"]
        self.client.post(f"/api/agents/{aid}/memory", json={"key": "k1", "content": "v1"})
        text, ok = app_module._run_readonly_sql("SELECT content FROM agent_memory WHERE agent_id = " + str(aid))
        self.assertTrue(ok)
        self.assertIn("v1", text)


class TestQaAazazAgent(QaBase):
    """Aazaz Ahmed: file-engine ops, path safety, agent_audit QC + download link."""

    def setUp(self):
        super().setUp()
        self._set("agent_enabled", "1")
        self._set("live_chat_ai", "0")
        self._set("review_enabled", "1")
        conn = sqlite3.connect(app_module.DB_PATH)
        conn.executescript("""
        INSERT INTO chat_agents (name, description, system_prompt, is_active, created_at, icon) VALUES
        ('Rumman Lashari', 'Administrator & Agent Coordinator - poore system ka boss.',
         'Aap Administrator hain.', 1, '2026-01-01 00:00:00', '');
        """)
        conn.commit()
        conn.close()

    def _aazaz_row(self):
        conn = sqlite3.connect(app_module.DB_PATH)
        conn.row_factory = sqlite3.Row
        row = conn.execute("SELECT * FROM chat_agents WHERE name = ?", (app_module._AAZAZ_NAME,)).fetchone()
        conn.close()
        return row

    def _seed_aazaz(self):
        conn = sqlite3.connect(app_module.DB_PATH)
        app_module._seed_aazaz(conn)
        conn.commit()
        conn.close()
        return self._aazaz_row()

    def test_seed_aazaz_idempotent(self):
        conn = sqlite3.connect(app_module.DB_PATH)
        app_module._seed_aazaz(conn)
        conn.commit()
        first = conn.execute(
            "SELECT COUNT(*) c FROM chat_agents WHERE name = ?", (app_module._AAZAZ_NAME,)
        ).fetchone()[0]
        mem1 = conn.execute(
            "SELECT COUNT(*) c FROM agent_memory WHERE agent_id = (SELECT id FROM chat_agents WHERE name = ?)",
            (app_module._AAZAZ_NAME,),
        ).fetchone()[0]
        app_module._seed_aazaz(conn)
        conn.commit()
        second = conn.execute(
            "SELECT COUNT(*) c FROM chat_agents WHERE name = ?", (app_module._AAZAZ_NAME,)
        ).fetchone()[0]
        mem2 = conn.execute(
            "SELECT COUNT(*) c FROM agent_memory WHERE agent_id = (SELECT id FROM chat_agents WHERE name = ?)",
            (app_module._AAZAZ_NAME,),
        ).fetchone()[0]
        conn.close()
        self.assertEqual(first, 1)
        self.assertEqual(second, 1)          # idempotent row
        self.assertGreaterEqual(mem1, 4)      # baseline memory seeded
        self.assertEqual(mem1, mem2)          # idempotent memory
        row = self._aazaz_row()
        self.assertEqual(row["icon"], "lucide:briefcase-business")

    def test_sanitize_fs_path_validation(self):
        with self.assertRaises(ValueError):
            app_module._sanitize_fs_path("relative/file.txt")
        with self.assertRaises(ValueError):
            app_module._sanitize_fs_path("C:\\Work\\..\\secret.txt")
        with self.assertRaises(ValueError):
            app_module._sanitize_fs_path("C:\\Work\\bad\x00name.txt")
        with self.assertRaises(ValueError):
            app_module._sanitize_fs_path("C:\\Work\\bad<name>.txt")
        with self.assertRaises(ValueError):
            app_module._sanitize_fs_path("")
        good = os.path.join(tempfile.mkdtemp(prefix="aazaz-path-"), "file.xlsx")
        self.assertEqual(app_module._sanitize_fs_path(good), os.path.abspath(good))

    def test_file_create_txt_read_and_delete_gate(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-txt-"), "daily-log.txt")
        text = app_module._run_agent_action(
            {"action": "file", "kind": "file", "file": {"op": "create", "path": tmp, "content": "Day 1: all good"}}
        )
        self.assertIn("**text file ready**", text)
        self.assertIn(tmp, text)
        self.assertIn("__filebadge__", text)
        with open(tmp, "r", encoding="utf-8") as fh:
            self.assertEqual(fh.read(), "Day 1: all good")
        # read
        out = app_module._run_file_action({"file": {"op": "read", "path": tmp}})
        self.assertIn("Day 1", out)
        # create again without overwrite -> refused
        with self.assertRaises(ValueError):
            app_module._run_agent_action({"action": "file", "kind": "file", "file": {"op": "create", "path": tmp, "content": "x"}})
        # overwrite explicitly allowed
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {"op": "create", "path": tmp, "content": "replaced", "overwrite": True}})
        with open(tmp, "r", encoding="utf-8") as fh:
            self.assertEqual(fh.read(), "replaced")
        # delete without confirmation -> refused
        with self.assertRaises(ValueError):
            app_module._run_file_action({"file": {"op": "delete", "path": tmp}})
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {"op": "delete", "path": tmp, "overwrite": True}})
        self.assertFalse(os.path.exists(tmp))

    def test_deletable_denies_system_areas(self):
        self.assertFalse(app_module._deletable("C:\\Windows\\System32\\evil.exe"))
        self.assertFalse(app_module._deletable("C:\\Users\\Me\\.ssh\\id_rsa"))
        self.assertFalse(app_module._deletable("C:\\x\\app.db"))
        self.assertTrue(app_module._deletable("C:\\Users\\Me\\Desktop\\report2.xlsx"))

    def test_file_xlsx_roundtrip(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-xlsx-"), "patients.xlsx")
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {
            "op": "create", "path": tmp,
            "sheet": "Patients",
            "header": ["Patient", "CPT", "Status"],
            "rows": [["Ali", "99213", "Pending"], ["Sana", "99214", "Done"]],
        }})
        from openpyxl import load_workbook
        wb = load_workbook(tmp)
        ws = wb.active
        self.assertEqual(ws.cell(1, 1).value, "Patient")
        self.assertEqual(ws.max_row, 3)
        self.assertEqual(ws.cell(3, 2).value, "99214")
        wb.close()
        out = app_module._run_file_action({"file": {"op": "read", "path": tmp}})
        self.assertIn("Patients", out)
        self.assertIn("99214", out)

    def test_file_docx_roundtrip(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-docx-"), "letter.docx")
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {
            "op": "create", "path": tmp,
            "content": "# Subject: remit note\n\nSir, yeh report ready hai.",
        }})
        self.assertTrue(os.path.getsize(tmp) > 0)
        from docx import Document
        doc = Document(tmp)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("Subject: remit note" in t for t in texts))
        self.assertTrue(any("report ready hai" in t for t in texts))
        out = app_module._run_file_action({"file": {"op": "read", "path": tmp}})
        self.assertIn("report ready hai", out)

    def test_file_pdf_create_and_read(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-pdf-"), "report.pdf")
        app_module._run_file_action({"file": {"op": "create", "path": tmp, "content": "# QC Report\n\nSab agents PASS."}})
        self.assertTrue(os.path.getsize(tmp) > 0)
        out = app_module._run_file_action({"file": {"op": "read", "path": tmp}})
        self.assertIn("QC Report", out)

    def test_xlsx_exec_formulas_formats_freeze_conditional(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-xlsx2-"), "ledger.xlsx")
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {
            "op": "create", "path": tmp,
            "sheet": "Ledger",
            "header": ["Item", "Qty", "Rate", "Total", "Share"],
            "rows": [["A", 3, 10, "=C2*B2", "=D2/$D$4"], ["B", 5, 8, "=C3*B3", "=D3/$D$4"],
                     ["Total", "", "", "=SUM(D2:D3)", ""]],
            "formats": {"3": "currency", "4": "currency", "2": "int", "5": "percent"},
            "conditional": {"2": {"op": ">", "value": 4, "fill": "FEE2E2"}},
            "color_scale": "4",
        }})
        from openpyxl import load_workbook
        wb = load_workbook(tmp)
        ws = wb["Ledger"]
        self.assertEqual(ws.freeze_panes, "A2")                    # frozen header
        self.assertEqual(ws["D2"].value, "=C2*B2")                 # formula injected
        self.assertEqual(ws["D4"].value, "=SUM(D2:D3)")            # nested logic
        self.assertEqual(ws["C2"].number_format, '"$"#,##0.00')    # currency column
        self.assertEqual(ws["B2"].number_format, "#,##0")          # integer column
        self.assertEqual(ws["E2"].number_format, "0.0%")           # percent column
        self.assertGreater(ws.column_dimensions["A"].width, 3)     # autofit applied
        self.assertGreaterEqual(len(list(ws.conditional_formatting)), 2)
        wb.close()

    def test_docx_exec_styles_and_callout(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-docx2-"), "exec.docx")
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {
            "op": "create", "path": tmp, "title": "Revenue Report",
            "content": "# Quarterly Report\n\n> Note: Revenue up 12% QoQ.\n\n"
                       "## Revenue\n\nWeek | Amount\n---|---\nW1 | $1,200.50\nW2 | $1,340.00\n",
        }})
        from docx import Document
        doc = Document(tmp)
        self.assertEqual(doc.paragraphs[0].text, "Revenue Report")         # Title
        self.assertIn("Generated:", doc.paragraphs[1].text)                # meta line
        self.assertEqual(doc.paragraphs[2].style.name, "Heading 1")         # H1 hierarchy
        self.assertGreaterEqual(len(doc.tables), 2)                        # callout + data table
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("Quarterly Report" in t for t in texts))
        self.assertTrue(any("Revenue" in t for t in texts))
        data_tbl = doc.tables[1]
        self.assertEqual(data_tbl.cell(0, 0).text, "Week")                 # styled header row
        self.assertEqual(data_tbl.cell(1, 1).text, "$1,200.50")

    def test_pdf_exec_header_footer_and_table(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-pdf2-"), "exec.pdf")
        app_module._run_agent_action({"action": "file", "kind": "file", "file": {
            "op": "create", "path": tmp, "title": "Ops Scorecard",
            "content": "# QC Report\n\n> Summary: 3 agents audited.\n\n"
                       "| Agent | Verdict |\n|---|---|\n| Aazaz Ahmed | PASS |\n| Medical Billing | WARNING |\n",
        }})
        self.assertTrue(os.path.getsize(tmp) > 1000)
        out = app_module._run_file_action({"file": {"op": "read", "path": tmp}})
        self.assertIn("QC Report", out)          # heading rendered
        self.assertIn("Aazaz Ahmed", out)        # zebra table row text

    def test_seed_memory_includes_advisory(self):
        self._seed_aazaz()
        conn = sqlite3.connect(app_module.DB_PATH)
        rows = conn.execute(
            "SELECT kind, key FROM agent_memory WHERE agent_id = (SELECT id FROM chat_agents WHERE name = ?)",
            (app_module._AAZAZ_NAME,),
        ).fetchall()
        conn.close()
        keys = {k for _, k in rows}
        self.assertTrue({"consultative", "formula_hint", "best_practice", "style_engine"} <= keys)

    def test_audit_logged_on_actions(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-audit-"), "n.txt")
        app_module._run_agent_action(
            {"action": "file", "kind": "file", "file": {"op": "create", "path": tmp, "content": "hello"}},
            agent_name="Aazaz Ahmed",
        )
        app_module._run_agent_action(
            {"action": "sql", "kind": "sql", "query": "SELECT COUNT(*) FROM tasks"},
            agent_name="Aazaz Ahmed",
        )
        with self.assertRaises(ValueError):
            app_module._run_agent_action(
                {"action": "file", "kind": "file", "file": {"op": "create", "path": tmp, "content": "x"}},
                agent_name="Aazaz Ahmed",
            )
        conn = sqlite3.connect(app_module.DB_PATH)
        rows = conn.execute("SELECT action, status FROM agent_audit ORDER BY id").fetchall()
        conn.close()
        self.assertEqual([tuple(r) for r in rows],
                         [("file", "ok"), ("sql", "ok"), ("file", "error")])

    def test_qc_report_aggregation(self):
        conn = sqlite3.connect(app_module.DB_PATH)
        app_module._seed_aazaz(conn)
        conn.commit()
        conn.close()
        app_module._audit_entry("Aazaz Ahmed", "file", "create", "/tmp/a.txt", "ok")
        app_module._audit_entry("Aazaz Ahmed", "file", "read", "/tmp/a.txt", "ok")
        app_module._audit_entry("Medical Billing", "sql", "sql", "SELECT ...", "error", "bad cpt")
        report = app_module._run_qc_report()
        self.assertIn("# System QC Audit", report)
        self.assertIn("Aazaz Ahmed", report)
        self.assertIn("PASS", report)
        self.assertIn("Recommendations", report)
        self.assertIn("bad cpt", report)

    def test_download_endpoint(self):
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-dl-"), "invoice.txt")
        with open(tmp, "w", encoding="utf-8") as fh:
            fh.write("INV-2026 amount 5000")
        p = urllib.parse.quote(tmp, safe="")
        r = self.client.get(f"/api/agents/files/download?p={p}")
        self.assertEqual(r.status_code, 200)
        self.assertIn(b"INV-2026", r.data)
        # unauthenticated client -> no session
        anon = self.app.test_client()
        r2 = anon.get(f"/api/agents/files/download?p={p}")
        self.assertEqual(r2.status_code, 401)
        # invalid path
        r3 = self.client.get("/api/agents/files/download?p=relative.txt")
        self.assertEqual(r3.status_code, 400)

    def test_normalize_file_plan(self):
        pl = app_module._normalize_agent_plan(
            {"action": "file", "file": {"op": "create", "path": "C:\\x\\y.xlsx", "content": "a", "overwrite": True}}
        )
        self.assertEqual(pl["action"], "file")
        self.assertEqual(pl["kind"], "file")
        self.assertEqual(pl["file"]["path"], "C:\\x\\y.xlsx")
        self.assertEqual(pl["file"]["op"], "create")
        self.assertIsNone(app_module._normalize_agent_plan({"action": "file", "file": {"op": "teleport"}}))
        self.assertIsNone(app_module._normalize_agent_plan({"action": "file"}))

    def test_qc_routing_hooks(self):
        aazaz = self._aazaz_row()
        if aazaz is None:
            aazaz = self._seed_aazaz()
        others = [
            {"id": 1, "name": "Rumman Lashari", "description": "Administrator"},
            {"id": aazaz["id"], "name": aazaz["name"], "description": aazaz["description"]},
        ]
        self.assertEqual(app_module._qc_agent(others)["name"], "Aazaz Ahmed")
        self.assertTrue(app_module._qc_file_intent("Excel bana do patients ka"))
        self.assertTrue(app_module._qc_file_intent("system audit karo"))
        self.assertFalse(app_module._qc_file_intent("good morning"))

    def test_file_plan_via_chat_maker_checked(self):
        self._seed_aazaz()
        sid = self._fresh_session()
        os.environ["GEMINI_API_KEY"] = "AIza-test"
        tmp = os.path.join(tempfile.mkdtemp(prefix="aazaz-chat-"), "from-chat.txt")
        plan = {"action": "file", "kind": "file", "id": None, "title": "", "fields": {},
                "file": {"op": "create", "path": tmp, "content": "chat se bani"}}
        try:
            with mock.patch.object(app_module, "_agent_plan", return_value=plan), \
                 mock.patch.object(app_module, "_review_draft", return_value=("approved", "")), \
                 mock.patch.object(app_module, "_llm_prompt", return_value=""):
                text, source, outcome = app_module.agent_answer(sid, "file bana do", "Tester", False)
            self.assertEqual(outcome, "action")
        finally:
            os.environ.pop("GEMINI_API_KEY", None)
        self.assertTrue(os.path.isfile(tmp))
        with open(tmp, "r", encoding="utf-8") as fh:
            self.assertEqual(fh.read(), "chat se bani")
        conn = sqlite3.connect(app_module.DB_PATH)
        n = conn.execute("SELECT COUNT(*) c FROM agent_audit WHERE action = 'file' AND status = 'ok'").fetchone()[0]
        conn.close()
        self.assertEqual(n, 1)


if __name__ == "__main__":
    unittest.main(verbosity=2)